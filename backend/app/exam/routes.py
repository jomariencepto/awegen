from flask import Blueprint, request, jsonify
from flask_jwt_extended import jwt_required, get_jwt_identity
import os
import io
from app.exam.service import ExamService
from app.exam.models import Exam, ExamQuestion
from app.utils.decorators import role_required
from app.auth.models import User
from app.utils.logger import get_logger

logger = get_logger(__name__)
PDF_DOWNLOAD_PASSWORD = os.getenv("PDF_DOWNLOAD_PASSWORD", "PDMEXAM@123")

exam_bp = Blueprint("exams", __name__)

# =========================
# Helper Functions
# =========================


def _get_dominant_difficulty(topic, questions_data):
    """Helper function to get dominant difficulty for a topic"""
    topic_questions = [q for q in questions_data if q.get('topic') == topic]
    if not topic_questions:
        return 'medium'
    
    difficulty_count = {}
    for q in topic_questions:
        diff = q.get('difficulty_level', 'medium')
        difficulty_count[diff] = difficulty_count.get(diff, 0) + 1
    
    return max(difficulty_count, key=difficulty_count.get) if difficulty_count else 'medium'


def _encrypt_pdf_buffer_if_needed(buffer):
    """Encrypt an in-memory PDF buffer using the shared password."""
    if not PDF_DOWNLOAD_PASSWORD:
        buffer.seek(0)
        return buffer
    try:
        from PyPDF2 import PdfReader, PdfWriter
        reader = PdfReader(buffer)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        writer.encrypt(PDF_DOWNLOAD_PASSWORD)
        out = io.BytesIO()
        writer.write(out)
        out.seek(0)
        return out
    except Exception as e:
        logger.warning(f"PDF encryption skipped: {e}")
        buffer.seek(0)
        return buffer


# =========================
# Exams CRUD
# =========================

@exam_bp.route("/all", methods=["GET"])
@jwt_required()
@role_required(["admin", "teacher", "department"])
def get_all_exams():
    """Get all exams (Admin Dashboard)"""
    try:
        page = request.args.get("page", 1, type=int)
        per_page = request.args.get("per_page", 10, type=int)
        status = request.args.get("status", None, type=str)
        
        result, status_code = ExamService.get_all_exams(page, per_page, status)
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in get_all_exams: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exams"
        }), 500


@exam_bp.route("", methods=["POST"])
@jwt_required()
@role_required(["teacher", "admin"])
def create_exam():
    """Create a new exam"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        # Convert string ID to integer for service
        teacher_id = int(get_jwt_identity())
        
        logger.info(f"Creating exam for teacher {teacher_id}")
        
        result, status_code = ExamService.create_exam(data, teacher_id)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in create_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to create exam",
            "error": str(e)
        }), 500


@exam_bp.route("/<int:exam_id>", methods=["GET"])
@jwt_required()
def get_exam(exam_id):
    """Get exam by ID"""
    try:
        result, status_code = ExamService.get_exam_by_id(exam_id)
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in get_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exam"
        }), 500


@exam_bp.route("/<int:exam_id>", methods=["PUT"])
@jwt_required()
@role_required(["teacher", "admin"])
def update_exam(exam_id):
    """Update an exam"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        result, status_code = ExamService.update_exam(exam_id, data)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in update_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to update exam"
        }), 500


@exam_bp.route("/<int:exam_id>", methods=["DELETE"])
@jwt_required()
@role_required(["teacher", "admin"])
def delete_exam(exam_id):
    """Delete an exam"""
    try:
        result, status_code = ExamService.delete_exam(exam_id)
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in delete_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to delete exam"
        }), 500


# =========================
# Save Exam Endpoint
# =========================

@exam_bp.route("/<int:exam_id>/save", methods=["PUT", "POST"])
@jwt_required()
@role_required(["teacher", "admin", "department", "department_head"])
def save_exam(exam_id):
    """Save exam as draft or update existing exam"""
    try:
        data = request.get_json()
        
        if data is None:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        teacher_id = int(get_jwt_identity())
        
        logger.info(f"Saving exam {exam_id} for teacher {teacher_id}")
        
        # Verify the exam belongs to the teacher
        exam = Exam.query.get(exam_id)
        
        if not exam:
            return jsonify({
                "success": False,
                "message": "Exam not found"
            }), 404
        
        # Check if user is the owner or admin
        current_user = User.query.get(teacher_id)
        if exam.teacher_id != teacher_id and current_user.role.lower() != "admin":
            return jsonify({
                "success": False,
                "message": "Unauthorized to save this exam"
            }), 403
        
        # Update the exam
        result, status_code = ExamService.update_exam(exam_id, data)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in save_exam: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": "Failed to save exam",
            "error": str(e)
        }), 500


# =========================
# Teacher Exams
# =========================

@exam_bp.route("/teacher/<int:teacher_id>", methods=["GET"])
@jwt_required()
def get_exams_by_teacher(teacher_id):
    """Get all exams by teacher"""
    try:
        current_user_id = int(get_jwt_identity())
        current_user = User.query.get(current_user_id)

        if not current_user:
            return jsonify({
                "success": False, 
                "message": "User not found"
            }), 404

        # Check authorization
        if current_user_id != teacher_id and current_user.role.lower() != "admin":
            return jsonify({
                "success": False, 
                "message": "Unauthorized"
            }), 403

        # Get pagination parameters
        page = request.args.get("page", 1, type=int)
        per_page = request.args.get("per_page", 10, type=int)

        result, status_code = ExamService.get_exams_by_teacher(
            teacher_id, page, per_page
        )
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in get_exams_by_teacher: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exams"
        }), 500


@exam_bp.route("/teacher/<int:teacher_id>/summary", methods=["GET"])
@jwt_required()
def get_teacher_dashboard_summary(teacher_id):
    """Get lightweight dashboard stats and recent exams for a teacher."""
    try:
        current_user_id = int(get_jwt_identity())
        current_user = User.query.get(current_user_id)

        if not current_user:
            return jsonify({
                "success": False,
                "message": "User not found"
            }), 404

        if current_user_id != teacher_id and current_user.role.lower() != "admin":
            return jsonify({
                "success": False,
                "message": "Unauthorized"
            }), 403

        result, status_code = ExamService.get_teacher_dashboard_summary(teacher_id)
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"Error in get_teacher_dashboard_summary: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get dashboard summary"
        }), 500
        
        
# =========================
# Saved / Draft Exams
# =========================

@exam_bp.route("/saved-exams", methods=["GET"])
@jwt_required()
@role_required(["teacher", "admin"])
def get_saved_exams():
    """Get saved (draft) exams for the current teacher"""
    try:
        teacher_id = int(get_jwt_identity())
        result, status_code = ExamService.get_saved_exams(teacher_id)
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"Error in get_saved_exams: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get saved exams"
        }), 500
        
    
# =========================
# Exam Workflow - SUBMIT ROUTE ⭐
# =========================

@exam_bp.route('/<int:exam_id>/submit', methods=['POST'])
@jwt_required()
def submit_exam_for_approval(exam_id):
    """
    ⭐ SUBMIT EXAM FOR APPROVAL
    Frontend: ManageExams.jsx
    URL: POST /api/exams/{exam_id}/submit
    """
    try:
        data = request.get_json()
        teacher_id = int(get_jwt_identity())  # ⭐ Convert to int
        
        logger.info(f"📤 Submit exam request: exam_id={exam_id}, teacher_id={teacher_id}")
        logger.info(f"📋 Submit data: {data}")
        
        # Verify exam exists
        exam = Exam.query.get(exam_id)
        if not exam:
            logger.error(f"❌ Exam {exam_id} not found")
            return jsonify({
                'success': False, 
                'message': 'Exam not found'
            }), 404
        
        # Verify ownership
        if exam.teacher_id != teacher_id:
            logger.error(f"❌ Teacher {teacher_id} doesn't own exam {exam_id} (owner: {exam.teacher_id})")
            return jsonify({
                'success': False, 
                'message': 'Unauthorized - You do not own this exam'
            }), 403
        
        # Add exam_id to data if not present
        if 'exam_id' not in data:
            data['exam_id'] = exam_id
        
        # Call service
        logger.info(f"✅ Calling ExamService.submit_exam_for_approval")
        result, status_code = ExamService.submit_exam_for_approval(data)
        
        logger.info(f"✅ Submit result: success={result.get('success')}, status={status_code}")
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"❌ Error in submit exam route: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False, 
            'message': f'Failed to submit exam: {str(e)}'
        }), 500


@exam_bp.route('/<int:exam_id>/reuse', methods=['POST'])
@jwt_required()
@role_required(['teacher', 'admin'])
def reuse_exam(exam_id):
    """
    Re-use an approved exam after 3 years.
    Creates a new draft copy of the exam with all questions.
    Frontend: ManageExams.jsx
    URL: POST /api/exams/{exam_id}/reuse
    """
    try:
        teacher_id = int(get_jwt_identity())
        logger.info(f"Re-use exam request: exam_id={exam_id}, teacher_id={teacher_id}")

        result, status_code = ExamService.reuse_exam(exam_id, teacher_id)
        return jsonify(result), status_code

    except Exception as e:
        logger.error(f"Error in reuse_exam: {str(e)}")
        return jsonify({
            'success': False,
            'message': f'Failed to re-use exam: {str(e)}'
        }), 500


@exam_bp.route("/approve", methods=["POST"])
@jwt_required()
@role_required(["admin", "department"])
def approve_exam():
    """Approve or reject an exam"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        # FIX: Pass admin_id to the service method
        admin_id = int(get_jwt_identity())
        result, status_code = ExamService.approve_exam(data, admin_id)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in approve_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to approve exam"
        }), 500


@exam_bp.route("/send-to-department", methods=["POST"])
@jwt_required()
@role_required(["admin", "teacher", "department_head"])
def send_exam_to_department():
    """Send exam to department (admins or the owning teacher)."""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        sender_id = int(get_jwt_identity())
        result, status_code = ExamService.send_exam_to_department(data, sender_id)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in send_exam_to_department: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to send exam to department"
        }), 500


# =========================
# Exam Generation
# =========================

@exam_bp.route("/generate-exam", methods=["POST"])
@jwt_required()
@role_required(["teacher", "admin"])
def generate_exam():
    try:
        data = request.get_json()
        
        # FORCE humanization
        if 'use_humanized' not in data:
            data['use_humanized'] = True
        
        teacher_id = int(get_jwt_identity())
        result, status_code = ExamService.create_exam(data, teacher_id)
        return jsonify(result), status_code
    
    except Exception as e:
        logger.error(f"Error in generate_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to generate exam",
            "error": str(e)
        }), 500


@exam_bp.route("/preview-exam", methods=["POST"])
@jwt_required()
@role_required(["teacher", "admin"])
def preview_exam():
    """Preview an exam without saving it"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        teacher_id = int(get_jwt_identity())
        
        logger.info(f"Previewing exam for teacher {teacher_id}")
        
        from app.exam.service import _get_exam_generator
        from app.module_processor.saved_module import SavedModuleService
        
        # Get content from all modules
        all_module_content = []
        for module_info in data['modules']:
            module_id = module_info['module_id']
            teaching_hours = module_info['teaching_hours']
            
            module_content_result, _ = SavedModuleService.get_module_content(module_id)
            if not module_content_result['success']:
                return jsonify({
                    'success': False, 
                    'message': f'Failed to get content for module {module_id}'
                }), 400
            
            for content_item in module_content_result['contents']:
                content_item['teaching_hours'] = teaching_hours
                content_item['module_id'] = module_id
            
            all_module_content.extend(module_content_result['contents'])
        
        # Prepare exam config
        exam_config = {
            'title': data['title'],
            'description': data.get('description', ''),
            'num_questions': data['num_questions'],
            'question_types': data['question_types'],
            'question_types_with_points': data.get('question_types_with_points', []),
            'cognitive_distribution': data.get('cognitive_distribution'),
            'duration_minutes': data['duration_minutes'],
            'passing_score': data['passing_score'],
            'total_hours': data['total_hours']
        }
        
        # Generate exam (cached per-process)
        exam_generator = _get_exam_generator()
        exam_result = exam_generator.generate_exam(all_module_content, exam_config)
        
        if not exam_result['success']:
            status_code = 400 if exam_result.get('error_code') == 'SCORE_TARGET_MISMATCH' else 500
            return jsonify({
                'success': False, 
                'message': exam_result.get('message', 'Failed to generate exam'),
                'error_code': exam_result.get('error_code'),
                'generated_points': exam_result.get('generated_points'),
                'target_points': exam_result.get('target_points'),
            }), status_code
        
        # Normalize question text for displays
        from app.exam.service import ExamService
        for question in exam_result.get('questions', []):
            if 'question_text' in question:
                question['question_text'] = ExamService._normalize_question_text_for_client(
                    question['question_text']
                )
        
        return jsonify({
            'success': True,
            'message': 'Exam preview generated successfully',
            'questions': exam_result['questions'],
            'tos': exam_result.get('tos', {}),
            'total_questions': len(exam_result['questions'])
        }), 200
        
    except Exception as e:
        logger.error(f"Error in preview_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to preview exam",
            "error": str(e)
        }), 500
        
       
@exam_bp.route("/preview/<int:exam_id>", methods=["GET"])
@jwt_required()
@role_required(["teacher", "admin"])
def preview_saved_exam(exam_id):
    """Preview an already saved exam"""
    try:
        import json
        
        logger.info(f"Previewing saved exam {exam_id}")
        
        # Get exam
        exam = Exam.query.get(exam_id)
        if not exam:
            return jsonify({
                "success": False,
                "message": "Exam not found"
            }), 404
        
        # Verify access (teacher who created it or admin)
        teacher_id = int(get_jwt_identity())
        current_user = User.query.get(teacher_id)
        
        if exam.teacher_id != teacher_id and current_user.role.lower() != "admin":
            return jsonify({
                "success": False,
                "message": "Unauthorized to preview this exam"
            }), 403
        
        # Get questions
        questions = ExamQuestion.query.filter_by(exam_id=exam_id).order_by(ExamQuestion.question_id).all()
        
        # Format questions for frontend
        from app.exam.service import ExamService
        formatted_questions = []
        for q in questions:
            question_dict = {
                "question_id": q.question_id,
                "question_text": ExamService._normalize_question_text_for_client(q.question_text),
                "question_type": q.question_type,
                "difficulty_level": q.difficulty_level,
                "correct_answer": q.correct_answer,
                "points": q.points or 1,
                "feedback": q.feedback,
                "image_id": q.image_id,
                "image_module_id": None
            }
            
            # Parse options if it's stored as JSON string
            if q.options:
                try:
                    if isinstance(q.options, str):
                        question_dict["options"] = json.loads(q.options)
                    else:
                        question_dict["options"] = q.options
                except:
                    question_dict["options"] = []
            else:
                question_dict["options"] = []
            
            # Map question_type to frontend display format
            type_mapping = {
                'multiple_choice': 'Multiple Choice',
                'true_false': 'True or False',
                'fill_in_blank': 'Fill in the Blanks',
                'identification': 'Identification',
                'matching_type': 'Matching Type',
                'short_answer': 'Short Answer',
                'essay': 'Essay'
            }
            question_dict["question_type"] = type_mapping.get(q.question_type, q.question_type.replace('_', ' ').title())
            
            # Resolve module_id for the image (needed by frontend to fetch it)
            if q.image_id:
                try:
                    from app.module_processor.models import ModuleImage
                    img = ModuleImage.query.get(q.image_id)
                    question_dict["image_module_id"] = img.module_id if img else None
                except Exception:
                    pass

            formatted_questions.append(question_dict)

        return jsonify({
            "success": True,
            "exam": {
                "exam_id": exam.exam_id,
                "title": exam.title,
                "description": exam.description,
                "duration_minutes": exam.duration_minutes,
                "passing_score": exam.passing_score,
                "total_questions": len(formatted_questions),
                "created_at": exam.created_at.isoformat() if exam.created_at else None
            },
            "questions": formatted_questions
        }), 200
        
    except Exception as e:
        logger.error(f"Error in preview_saved_exam: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Failed to preview exam: {str(e)}"
        }), 500


# =========================
# TOS Endpoint
# =========================

@exam_bp.route("/<int:exam_id>/tos", methods=["GET"])
@jwt_required()
def get_exam_tos(exam_id):
    """Get Table of Specification for an exam"""
    try:
        from app.exam.tos_generator import TOSGenerator
        from app.exam.bloom_classifier import BloomClassifier
        import json
        
        # Get exam
        exam = Exam.query.get(exam_id)
        if not exam:
            return jsonify({
                "success": False,
                "message": "Exam not found"
            }), 404
        
        # Get questions
        questions = ExamQuestion.query.filter_by(exam_id=exam_id).all()
        
        if not questions:
            return jsonify({
                "success": False,
                "message": "No questions found for this exam"
            }), 404
        
        # Convert questions to dict format
        questions_data = []
        for q in questions:
            question_dict = q.to_dict()
            # Parse options if stored as JSON string
            if isinstance(question_dict.get('options'), str):
                try:
                    question_dict['options'] = json.loads(question_dict['options'])
                except:
                    question_dict['options'] = []
            questions_data.append(question_dict)
        
        # Classify questions if they don't have bloom_level
        bloom_classifier = BloomClassifier()
        for q in questions_data:
            if 'bloom_level' not in q or not q.get('bloom_level'):
                q['bloom_level'] = bloom_classifier.classify_question(q['question_text'])
        
        # Extract topics from questions or use module
        topics = set()
        for q in questions_data:
            if q.get('topic'):
                topics.add(q['topic'])
        
        # If no topics in questions, use exam module title
        if not topics and exam.module:
            topics.add(exam.module.title)
        
        if not topics:
            topics.add('General')
        
        topics = list(topics)
        
        # Generate TOS
        tos_generator = TOSGenerator()
        exam_config = {
            'title': exam.title,
            'duration_minutes': exam.duration_minutes,
        }
        
        tos = tos_generator.generate_tos(questions_data, topics, exam_config)
        
        # Calculate difficulty distribution if not present
        if 'difficulty_distribution' not in tos or not tos['difficulty_distribution']:
            difficulty_dist = {}
            for q in questions_data:
                diff = q.get('difficulty_level', 'medium')
                difficulty_dist[diff] = difficulty_dist.get(diff, 0) + 1
            tos['difficulty_distribution'] = difficulty_dist
        
        # Format response to match frontend expectations
        response_data = {
            "success": True,
            "exam_title": exam.title,
            "total_questions": len(questions_data),
            "topics_count": len(topics),
            "cognitive_levels_count": len([k for k, v in tos['cognitive_distribution'].items() if v > 0]),
            
            # Cognitive levels for frontend
            "cognitive_levels": [
                {
                    "level": level.capitalize(),
                    "count": tos['cognitive_distribution'].get(level, 0),
                    "percentage": round(tos['cognitive_percentages'].get(level, 0), 2),
                    "target_percentage": round(tos['cognitive_percentages'].get(level, 0), 2)
                }
                for level in ['remembering', 'understanding', 'applying', 'analyzing', 'evaluating', 'creating']
                if tos['cognitive_distribution'].get(level, 0) > 0
            ],
            
            # Topics for frontend
            "topics": [
                {
                    "topic_id": idx,
                    "name": topic,
                    "question_count": sum(tos['topic_cognitive_matrix'].get(topic, {}).values()),
                    "percentage": round((sum(tos['topic_cognitive_matrix'].get(topic, {}).values()) / len(questions_data) * 100), 2) if questions_data else 0,
                    "difficulty": _get_dominant_difficulty(topic, questions_data)
                }
                for idx, topic in enumerate(topics)
            ],
            
            # Matrix data
            "matrix": [
                {
                    "topic_id": idx,
                    "topic_name": topic,
                    "distribution": [
                        tos['topic_cognitive_matrix'].get(topic, {}).get(level, 0)
                        for level in ['remembering', 'understanding', 'applying', 'analyzing', 'evaluating', 'creating']
                    ],
                    "total": sum(tos['topic_cognitive_matrix'].get(topic, {}).values())
                }
                for idx, topic in enumerate(topics)
            ],
            
            # Additional data
            "difficulty_distribution": tos.get('difficulty_distribution', {}),
            "question_type_distribution": tos.get('question_type_distribution', {}),
            "summary": tos.get('summary', {})
        }
        
        return jsonify(response_data), 200
        
    except Exception as e:
        logger.error(f"Error in get_exam_tos: {str(e)}")
        import traceback
        traceback.print_exc()
        return jsonify({
            "success": False,
            "message": f"Failed to get TOS: {str(e)}"
        }), 500


# =========================
# Categories & Questions
# =========================

@exam_bp.route("/categories", methods=["GET"])
@jwt_required()
def get_exam_categories():
    """Get all exam categories"""
    try:
        result, status_code = ExamService.get_exam_categories()
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in get_exam_categories: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exam categories"
        }), 500


@exam_bp.route("/<int:exam_id>/questions", methods=["GET"])
@jwt_required()
def get_exam_questions(exam_id):
    """Get all questions for an exam"""
    try:
        result, status_code = ExamService.get_exam_by_id(exam_id)
        if result.get('success'):
            return jsonify({
                "success": True,
                "questions": result.get('questions', [])
            }), 200
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in get_exam_questions: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exam questions"
        }), 500


@exam_bp.route("/<int:exam_id>/questions", methods=["POST"])
@jwt_required()
@role_required(["teacher", "admin"])
def add_question_to_exam(exam_id):
    """Add a question to an exam"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        result, status_code = ExamService.add_question_to_exam(exam_id, data)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in add_question_to_exam: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to add question to exam"
        }), 500


@exam_bp.route("/questions/<int:question_id>", methods=["GET"])
@jwt_required()
def get_question(question_id):
    """Get a specific question"""
    try:
        question = ExamQuestion.query.get(question_id)
        
        if not question:
            return jsonify({
                "success": False,
                "message": "Question not found"
            }), 404
        
        return jsonify({
            "success": True,
            "question": question.to_dict()
        }), 200
        
    except Exception as e:
        logger.error(f"Error in get_question: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get question"
        }), 500


@exam_bp.route("/questions/<int:question_id>", methods=["PUT"])
@jwt_required()
@role_required(["teacher", "admin"])
def update_question(question_id):
    """Update a question"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        result, status_code = ExamService.update_question(question_id, data)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in update_question: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to update question"
        }), 500


@exam_bp.route("/questions/<int:question_id>/feedback", methods=["PUT"])
@jwt_required()
@role_required(["department", "admin"])
def update_question_feedback(question_id):
    """Update feedback for a specific question"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        result, status_code = ExamService.update_question_feedback(question_id, data)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in update_question_feedback: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to update feedback"
        }), 500


@exam_bp.route("/questions/<int:question_id>", methods=["DELETE"])
@jwt_required()
@role_required(["teacher", "admin"])
def delete_question(question_id):
    """Delete a question"""
    try:
        result, status_code = ExamService.delete_question(question_id)
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in delete_question: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to delete question"
        }), 500


# =========================
# Submissions
# =========================

@exam_bp.route("/submissions", methods=["POST"])
@jwt_required()
def submit_exam_submission():
    """Submit exam answers"""
    try:
        data = request.get_json()
        
        if not data:
            return jsonify({
                "success": False,
                "message": "No data provided"
            }), 400
        
        user_id = int(get_jwt_identity())
        
        result, status_code = ExamService.submit_exam_submission(data, user_id)
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in submit_exam_submission: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to submit exam"
        }), 500


@exam_bp.route("/<int:exam_id>/submissions", methods=["GET"])
@jwt_required()
@role_required(["teacher", "admin"])
def get_exam_submissions(exam_id):
    """Get all submissions for an exam"""
    try:
        page = request.args.get("page", 1, type=int)
        per_page = request.args.get("per_page", 10, type=int)

        result, status_code = ExamService.get_exam_submissions(
            exam_id, page, per_page
        )
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in get_exam_submissions: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exam submissions"
        }), 500


@exam_bp.route("/user/<int:user_id>/submissions", methods=["GET"])
@jwt_required()
def get_user_submissions(user_id):
    """Get all submissions by a user"""
    try:
        current_user_id = int(get_jwt_identity())
        current_user = User.query.get(current_user_id)

        if not current_user:
            return jsonify({
                "success": False, 
                "message": "User not found"
            }, 404)

        # Check authorization
        if current_user_id != user_id and current_user.role.lower() != "admin":
            return jsonify({
                "success": False, 
                "message": "Unauthorized"
            }), 403

        page = request.args.get("page", 1, type=int)
        per_page = request.args.get("per_page", 10, type=int)

        result, status_code = ExamService.get_user_submissions(
            user_id, page, per_page
        )
        return jsonify(result), status_code
        
    except Exception as e:
        logger.error(f"Error in get_user_submissions: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get user submissions"
        }), 500


@exam_bp.route("/submissions/<int:submission_id>", methods=["GET"])
@jwt_required()
def get_submission(submission_id):
    """Get a specific submission"""
    try:
        from app.exam.models import ExamSubmission
        submission = ExamSubmission.query.get(submission_id)
        
        if not submission:
            return jsonify({
                "success": False,
                "message": "Submission not found"
            }), 404
        
        # Check authorization
        current_user_id = int(get_jwt_identity())
        current_user = User.query.get(current_user_id)
        
        if not current_user:
            return jsonify({
                "success": False,
                "message": "User not found"
            }, 404)
        
        # Allow access if user is the submitter, the exam creator, or an admin
        if (submission.user_id != current_user_id and 
            submission.exam.teacher_id != current_user_id and 
            current_user.role.lower() != "admin"):
            return jsonify({
                "success": False,
                "message": "Unauthorized"
            }), 403
        
        return jsonify({
            "success": True,
            "submission": submission.to_dict()
        }), 200
        
    except Exception as e:
        logger.error(f"Error in get_submission: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get submission"
        }), 500


@exam_bp.route("/submissions/<int:submission_id>/answers", methods=["GET"])
@jwt_required()
def get_submission_answers(submission_id):
    """Get all answers for a submission"""
    try:
        result, status_code = ExamService.get_submission_answers(submission_id)
        return jsonify(result), status_code
    except Exception as e:
        logger.error(f"Error in get_submission_answers: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get submission answers"
        }), 500


# =========================
# Analytics & Statistics
# =========================

@exam_bp.route("/<int:exam_id>/statistics", methods=["GET"])
@jwt_required()
@role_required(["teacher", "admin"])
def get_exam_statistics(exam_id):
    """Get statistics for an exam"""
    try:
        from app.exam.models import ExamSubmission
        from sqlalchemy import func
        
        exam = Exam.query.get(exam_id)
        if not exam:
            return jsonify({
                "success": False,
                "message": "Exam not found"
            }), 404
        
        # Calculate statistics
        submissions = ExamSubmission.query.filter_by(
            exam_id=exam_id,
            is_completed=True
        ).all()
        
        total_submissions = len(submissions)
        
        if total_submissions == 0:
            return jsonify({
                "success": True,
                "statistics": {
                    "total_submissions": 0,
                    "average_score": 0,
                    "highest_score": 0,
                    "lowest_score": 0,
                    "pass_rate": 0
                }
            }), 200
        
        scores = [s.score for s in submissions if s.score is not None]
        total_points = submissions[0].total_points if submissions else 0
        
        average_score = sum(scores) / len(scores) if scores else 0
        highest_score = max(scores) if scores else 0
        lowest_score = min(scores) if scores else 0
        
        # Calculate pass rate
        passing_score = exam.passing_score or 75
        passed = sum(1 for s in submissions 
                    if s.score and s.total_points and 
                    (s.score / s.total_points * 100) >= passing_score)
        pass_rate = (passed / total_submissions * 100) if total_submissions > 0 else 0
        
        return jsonify({
            "success": True,
            "statistics": {
                "total_submissions": total_submissions,
                "average_score": round(average_score, 2),
                "highest_score": highest_score,
                "lowest_score": lowest_score,
                "pass_rate": round(pass_rate, 2),
                "total_points": total_points
            }
        }), 200
        
    except Exception as e:
        logger.error(f"Error in get_exam_statistics: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get exam statistics"
        }), 500


@exam_bp.route("/<int:exam_id>/question-analysis", methods=["GET"])
@jwt_required()
@role_required(["teacher", "admin"])
def get_question_analysis(exam_id):
    """Get question-level analysis for an exam"""
    try:
        from app.exam.models import ExamAnswer, ExamSubmission
        from sqlalchemy import func
        
        exam = Exam.query.get(exam_id)
        if not exam:
            return jsonify({
                "success": False,
                "message": "Exam not found"
            }), 404
        
        # Get all questions for the exam
        questions = ExamQuestion.query.filter_by(exam_id=exam_id).all()
        
        question_analysis = []
        
        for question in questions:
            # Get all answers for this question
            answers = ExamAnswer.query.join(ExamSubmission).filter(
                ExamAnswer.question_id == question.question_id,
                ExamSubmission.exam_id == exam_id,
                ExamSubmission.is_completed == True
            ).all()
            
            total_answers = len(answers)
            correct_answers = sum(1 for a in answers if a.is_correct)
            
            # Calculate success rate
            success_rate = (correct_answers / total_answers * 100) if total_answers > 0 else 0
            
            # Normalize question text for display
            question_text = question.question_text
            if len(question_text) > 100:
                question_text = question_text[:100] + "..."
            question_text = ExamService._normalize_question_text_for_client(question_text)
            
            question_analysis.append({
                "question_id": question.question_id,
                "question_text": question_text,
                "question_type": question.question_type,
                "difficulty": question.difficulty_level,
                "total_answers": total_answers,
                "correct_answers": correct_answers,
                "success_rate": round(success_rate, 2)
            })
        
        return jsonify({
            "success": True,
            "question_analysis": question_analysis
        }), 200
        
    except Exception as e:
        logger.error(f"Error in get_question_analysis: {str(e)}")
        return jsonify({
            "success": False,
            "message": "Failed to get question analysis"
        }), 500


# =========================
# Health Check
# =========================

@exam_bp.route("/health", methods=["GET"])
def health_check():
    """Health check endpoint for exam service"""
    return jsonify({
        "success": True,
        "service": "exam",
        "status": "healthy"
    }), 200

# =========================
# ANSWER KEY EXPORT ROUTES
# =========================


@exam_bp.route('/exports/exam/<int:exam_id>/answer-key/pdf', methods=['GET'])
@jwt_required()
@role_required(['teacher', 'admin'])
def export_answer_key_pdf(exam_id):
    """Export exam answer key as PDF"""
    try:
        from flask import send_file
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        import io
        
        logger.info(f"Exporting answer key PDF for exam {exam_id}")
        
        # Get answer key data
        result, status_code = ExamService.get_answer_key(exam_id)
        
        if status_code != 200:
            return jsonify(result), status_code
        
        answer_key_data = result['answer_key']
        
        # Create PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter,
                              rightMargin=72, leftMargin=72,
                              topMargin=72, bottomMargin=18)
        
        # Container for the 'Flowable' objects
        elements = []
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a56db'),
            spaceAfter=30,
            alignment=1  # Center
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#1f2937'),
            spaceAfter=12,
            spaceBefore=12
        )
        
        # Title
        title = Paragraph(f"<b>ANSWER KEY</b><br/>{answer_key_data['title']}", title_style)
        elements.append(title)
        elements.append(Spacer(1, 0.3*inch))
        
        # Exam Info
        info_data = [
            ['Total Questions:', str(answer_key_data['total_questions'])],
            ['Total Points:', str(answer_key_data['total_points'])],
            ['Duration:', f"{answer_key_data['duration_minutes']} minutes"],
            ['Teacher:', answer_key_data['teacher_name']]
        ]
        
        info_table = Table(info_data, colWidths=[2*inch, 4*inch])
        info_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f3f4f6')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e5e7eb'))
        ]))
        
        elements.append(info_table)
        elements.append(Spacer(1, 0.5*inch))
        
        # Answer Key Section
        elements.append(Paragraph("<b>CORRECT ANSWERS</b>", heading_style))
        elements.append(Spacer(1, 0.2*inch))
        
        # Questions and Answers
        for i, qa in enumerate(answer_key_data['questions'], 1):
            # Question
            question_style = ParagraphStyle(
                'Question',
                parent=styles['Normal'],
                fontSize=11,
                textColor=colors.black,
                spaceAfter=6,
                leftIndent=0
            )
            
            question_text = f"<b>{i}. {qa['question_text']}</b>"
            elements.append(Paragraph(question_text, question_style))
            
            # Answer
            answer_style = ParagraphStyle(
                'Answer',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.HexColor('#059669'),
                spaceAfter=12,
                leftIndent=20
            )
            
            answer_text = f"<b>✓ Answer:</b> {qa['correct_answer']} <b>({qa['points']} point{'s' if qa['points'] != 1 else ''})</b>"
            elements.append(Paragraph(answer_text, answer_style))
            
            # Question Type and Difficulty
            meta_style = ParagraphStyle(
                'Meta',
                parent=styles['Normal'],
                fontSize=8,
                textColor=colors.HexColor('#6b7280'),
                spaceAfter=18,
                leftIndent=20
            )
            
            meta_text = f"Type: {qa['question_type']} | Difficulty: {qa['difficulty_level']}"
            elements.append(Paragraph(meta_text, meta_style))
        
        # Build PDF
        doc.build(elements)
        buffer = _encrypt_pdf_buffer_if_needed(buffer)
        
        # Send file
        return send_file(
            buffer,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f'answer_key_{answer_key_data["title"].replace(" ", "_").lower()}.pdf'
        )
        
    except Exception as e:
        logger.error(f"Error exporting answer key PDF: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Failed to export answer key: {str(e)}'
        }), 500


@exam_bp.route('/exports/exam/<int:exam_id>/answer-key/word', methods=['GET'])
@jwt_required()
@role_required(['teacher', 'admin'])
def export_answer_key_word(exam_id):
    """Export exam answer key as DOCX"""
    try:
        from flask import send_file
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        import io
        
        logger.info(f"Exporting answer key DOCX for exam {exam_id}")
        
        # Get answer key data
        result, status_code = ExamService.get_answer_key(exam_id)
        
        if status_code != 200:
            return jsonify(result), status_code
        
        answer_key_data = result['answer_key']
        
        # Create Word document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading('ANSWER KEY', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(26, 86, 219)
        
        # Exam Title
        exam_title = doc.add_heading(answer_key_data['title'], 1)
        exam_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()  # Spacer
        
        # Exam Info Table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Grid Accent 1'
        
        table.rows[0].cells[0].text = 'Total Questions:'
        table.rows[0].cells[1].text = str(answer_key_data['total_questions'])
        
        table.rows[1].cells[0].text = 'Total Points:'
        table.rows[1].cells[1].text = str(answer_key_data['total_points'])
        
        table.rows[2].cells[0].text = 'Duration:'
        table.rows[2].cells[1].text = f"{answer_key_data['duration_minutes']} minutes"
        
        table.rows[3].cells[0].text = 'Teacher:'
        table.rows[3].cells[1].text = answer_key_data['teacher_name']
        
        # Bold first column
        for row in table.rows:
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()  # Spacer
        
        # Answer Key Section
        heading = doc.add_heading('CORRECT ANSWERS', 2)
        heading_run = heading.runs[0]
        heading_run.font.color.rgb = RGBColor(31, 41, 55)
        
        doc.add_paragraph()  # Spacer
        
        # Questions and Answers
        for i, qa in enumerate(answer_key_data['questions'], 1):
            # Question
            question = doc.add_paragraph()
            question_run = question.add_run(f"{i}. {qa['question_text']}")
            question_run.bold = True
            question_run.font.size = Pt(11)
            
            # Answer
            answer = doc.add_paragraph()
            answer.paragraph_format.left_indent = Inches(0.3)
            
            answer_label = answer.add_run('✓ Answer: ')
            answer_label.bold = True
            answer_label.font.color.rgb = RGBColor(5, 150, 105)
            
            answer_text = answer.add_run(qa['correct_answer'])
            answer_text.font.color.rgb = RGBColor(5, 150, 105)
            
            points_text = answer.add_run(f" ({qa['points']} point{'s' if qa['points'] != 1 else ''})")
            points_text.bold = True
            points_text.font.color.rgb = RGBColor(5, 150, 105)
            
            # Metadata
            meta = doc.add_paragraph()
            meta.paragraph_format.left_indent = Inches(0.3)
            meta_run = meta.add_run(f"Type: {qa['question_type']} | Difficulty: {qa['difficulty_level']}")
            meta_run.font.size = Pt(8)
            meta_run.font.color.rgb = RGBColor(107, 114, 128)
            
            doc.add_paragraph()  # Spacer between questions
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Send file
        return send_file(
            buffer,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'answer_key_{answer_key_data["title"].replace(" ", "_").lower()}.docx'
        )
        
    except Exception as e:
        logger.error(f"Error exporting answer key DOCX: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        return jsonify({
            'success': False,
            'message': f'Failed to export answer key: {str(e)}'
        }), 500


# =========================
# Math Solver Endpoint
# =========================

@exam_bp.route("/math/solve", methods=["POST"])
@jwt_required()
def math_solve():
    """
    POST /api/exams/math/solve
    Body: { "expression": "2*x + 6 = 14" }
    Response: { "success": bool, "numeric_value": float|None, "latex": str|None, "error": str|None }
    """
    try:
        from app.exam.math_solver import try_sympy_solve
        data = request.get_json(force=True) or {}
        expr = data.get('expression', '').strip()
        if not expr:
            return jsonify({'success': False, 'error': 'expression is required'}), 400
        result = try_sympy_solve(expr)
        return jsonify(result), 200
    except Exception as e:
        logger.error(f"math_solve error: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
