import os
import json
import uuid
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from flask import current_app, has_app_context
from app.utils.logger import get_logger

logger = get_logger(__name__)

_PASSTHROUGH_IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.jfif'}


def _get_backend_root():
    if has_app_context():
        return os.path.dirname(current_app.root_path)
    return str(Path(__file__).resolve().parents[2])


def _get_app_root():
    if has_app_context():
        return current_app.root_path
    return os.path.join(_get_backend_root(), 'app')


def _get_export_image_temp_path(image_id, suffix='.png'):
    temp_dir = os.path.join(_get_backend_root(), 'temp', 'export_images')
    os.makedirs(temp_dir, exist_ok=True)
    return os.path.join(temp_dir, f"question_image_{image_id}_{uuid.uuid4().hex}{suffix}")


def _resolve_module_image_source_path(module_image):
    backend_dir = _get_backend_root()
    app_root = _get_app_root()
    candidates = []

    if module_image.image_path:
        candidates.append(module_image.image_path)

        if not os.path.isabs(module_image.image_path):
            candidates.append(os.path.join(backend_dir, module_image.image_path))
            candidates.append(os.path.join(app_root, module_image.image_path))

        if f"{os.sep}app{os.sep}" in module_image.image_path:
            stripped = module_image.image_path.replace(f"{os.sep}app{os.sep}", os.sep)
            candidates.append(stripped)
            candidates.append(os.path.join(backend_dir, stripped))

        for root in (backend_dir, app_root):
            for folder_name in ('module_images', 'modules_images'):
                base_dir = os.path.join(root, 'uploads', folder_name, str(module_image.module_id))
                candidates.append(os.path.join(base_dir, os.path.basename(module_image.image_path)))

    chosen = next((path for path in candidates if path and os.path.exists(path)), None)
    return os.path.abspath(chosen) if chosen else None


def _convert_image_to_png(source_path, image_id):
    try:
        from PIL import Image
    except ImportError:
        logger.warning("Pillow is unavailable; skipping non-PNG/JPEG export image conversion.")
        return None, []

    temp_path = _get_export_image_temp_path(image_id, '.png')

    try:
        with Image.open(source_path) as pil_image:
            detected_format = str(pil_image.format or "").upper()
            file_ext = os.path.splitext(source_path)[1].lower()

            if detected_format in {"WMF", "EMF"} or file_ext in {".wmf", ".emf"}:
                try:
                    pil_image.load(dpi=300)
                except TypeError:
                    pil_image.load()
                except Exception:
                    pass

            if pil_image.mode in ('RGBA', 'LA'):
                base = pil_image.convert('RGBA')
                background = Image.new('RGBA', base.size, (255, 255, 255, 255))
                rendered = Image.alpha_composite(background, base).convert('RGB')
            elif pil_image.mode != 'RGB':
                rendered = pil_image.convert('RGB')
            else:
                rendered = pil_image.copy()

            rendered.save(temp_path, format='PNG', optimize=True)
            return temp_path, [temp_path]
    except Exception as exc:
        logger.warning(
            f"Failed to convert question image {image_id} for export from {source_path}: {exc}"
        )
        try:
            if os.path.exists(temp_path):
                os.remove(temp_path)
        except OSError:
            pass
        return None, []


def _resolve_question_image_for_export(question):
    image_id = question.get('image_id')
    if not image_id:
        return None, []

    try:
        from app.module_processor.models import ModuleImage

        module_image = ModuleImage.query.get(image_id)
        if not module_image:
            logger.warning(f"Question image {image_id} was referenced during export but not found.")
            return None, []
    except Exception as exc:
        logger.warning(f"Unable to load question image {image_id} for export: {exc}")
        return None, []

    source_path = _resolve_module_image_source_path(module_image)
    if not source_path:
        logger.warning(
            f"Question image file missing on disk for image_id={image_id}, module_id={module_image.module_id}"
        )
        return None, []

    source_ext = os.path.splitext(source_path)[1].lower()
    if source_ext in _PASSTHROUGH_IMAGE_EXTENSIONS:
        return source_path, []

    return _convert_image_to_png(source_path, image_id)


class WordExporter:
    """
    Word Exporter - M-IT322-HCI-Sem2-TQ.docx EXACT FORMAT with COMPRESSED SPACING
    
    Handles ALL downloads with MINIMAL spacing for maximum content per page
    ✅ ALL existing functions preserved
    ✅ 100% format match with compressed spacing
    """
    
    def __init__(self, logo_path=None):
        """Initialize Word Exporter"""
        if logo_path:
            self.logo_path = logo_path
        else:
            # Use Path(__file__).resolve() so the path is always correct
            # regardless of the working directory at startup.
            self.logo_path = str(Path(__file__).resolve().parent / 'assets' / 'school_logo.jpeg')
        self._temp_export_paths = []
    
    def _sanitize_text(self, text):
        """Sanitize text for Word export"""
        if text is None:
            return ""
        return str(text).strip()

    def _cleanup_temp_export_paths(self):
        for path in self._temp_export_paths:
            try:
                if path and os.path.exists(path):
                    os.remove(path)
            except OSError as exc:
                logger.warning(f"Failed to clean temporary export image {path}: {exc}")
        self._temp_export_paths = []

    def _add_question_image(self, doc, question):
        image_path, cleanup_paths = _resolve_question_image_for_export(question)
        self._temp_export_paths.extend(cleanup_paths)

        if not image_path or not os.path.exists(image_path):
            return

        image_width = Inches(4.8)

        try:
            from PIL import Image as PILImage

            with PILImage.open(image_path) as pil_image:
                width_px, height_px = pil_image.size

            if width_px > 0 and height_px > 0:
                native_width = width_px / 96.0
                native_height = height_px / 96.0
                scale = min(4.8 / native_width, 3.6 / native_height, 2.0)
                image_width = Inches(native_width * scale)
        except Exception:
            pass

        try:
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.0
            paragraph.add_run().add_picture(image_path, width=image_width)
        except Exception as exc:
            logger.warning(f"Failed to embed question image in DOCX export: {exc}")
    
    def _add_exam_header(self, doc, exam_data=None):
        """
        Add header matching the reference docx exactly — NO TABLE, uses inline
        picture + tab stops so there are zero cell borders possible.

        Layout (single-line each row, tab-stop at right margin):
          [Logo] PAMBAYANG DALUBHASAAN NG MARILAO  →  COLLEGE OF COMPUTER STUDIES
                 Abangan Norte, Marilao, Bulacan   →  Information Technology Department
        Followed by a paragraph with a bottom border as the horizontal rule.
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Twips

        # Usable width in twips: page width 12240 - left margin 1080 - right margin 1080 = 10080
        # (0.75 in margins = 1080 twips each side)
        RIGHT_TAB = 9360  # twips from left margin to right edge (6.5 in)

        def _set_right_tab(paragraph, pos_twips):
            """Add a right-aligned tab stop at pos_twips."""
            pPr = paragraph._p.get_or_add_pPr()
            tabs = OxmlElement('w:tabs')
            tab = OxmlElement('w:tab')
            tab.set(qn('w:val'), 'right')
            tab.set(qn('w:pos'), str(pos_twips))
            tabs.append(tab)
            pPr.append(tabs)

        def _set_para_spacing(paragraph, before=0, after=0):
            paragraph.paragraph_format.space_before = Pt(before)
            paragraph.paragraph_format.space_after = Pt(after)

        # ── Row 1: Logo + School Name (left) | College Name (right) ──────
        p1 = doc.add_paragraph()
        _set_para_spacing(p1, before=0, after=0)
        _set_right_tab(p1, RIGHT_TAB)

        # Logo inline
        run_logo = p1.add_run()
        if os.path.exists(self.logo_path):
            run_logo.add_picture(self.logo_path, width=Inches(0.55))
        else:
            logger.warning(f"Logo not found at: {self.logo_path}")

        # School name (bold, left side)
        run_school = p1.add_run("  PAMBAYANG DALUBHASAAN NG MARILAO")
        run_school.font.bold = True
        run_school.font.size = Pt(10)

        # Tab → right side: College name
        run_tab1 = p1.add_run("\t")
        run_col = p1.add_run("COLLEGE OF COMPUTER STUDIES")
        run_col.font.bold = True
        run_col.font.size = Pt(10)

        # ── Row 2: Address (left) | Department (right) ────────────────────
        p2 = doc.add_paragraph()
        _set_para_spacing(p2, before=0, after=4)
        _set_right_tab(p2, RIGHT_TAB)

        # Indent address to align under school name (~logo width + 2 spaces)
        run_indent = p2.add_run("       Abangan Norte, Marilao, Bulacan")
        run_indent.font.size = Pt(9)

        # Tab → right side: Department
        p2.add_run("\t")
        run_dept = p2.add_run("Information Technology Department")
        run_dept.font.size = Pt(9)

        # ── Bottom border rule (replicate the line under the reference header) ─
        p_rule = doc.add_paragraph()
        _set_para_spacing(p_rule, before=0, after=6)
        pPr = p_rule._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
    
    def _get_exam_type_title(self, exam_data, is_answer_key=False, is_special=False):
        """Get exam type title based on category"""
        category_name = exam_data.get('category_name', exam_data.get('category', 'MIDTERM'))
        
        if 'midterm' in str(category_name).lower():
            exam_type = 'MIDTERM EXAMINATION'
        elif 'final' in str(category_name).lower():
            exam_type = 'FINAL EXAMINATION'
        elif 'prelim' in str(category_name).lower():
            exam_type = 'PRELIMINARY EXAMINATION'
        else:
            exam_type = f'{str(category_name).upper()} EXAMINATION'
        
        if is_special:
            exam_type = f'SPECIAL {exam_type}'
        if is_answer_key:
            exam_type = f'{exam_type} - ANSWER KEY'
        
        return exam_type
    
    def _add_exam_title(self, doc, exam_data, is_answer_key=False, is_special=False):
        """
        Add header identity row + exam title matching department download format:
          Created by: <Teacher>    <EXAM TYPE>    Subject: <Subject>
                                <EXAM TITLE>
        """
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        def _set_center_right_tabs(paragraph, center_twips=4680, right_twips=9360):
            pPr = paragraph._p.get_or_add_pPr()
            tabs = OxmlElement('w:tabs')

            center = OxmlElement('w:tab')
            center.set(qn('w:val'), 'center')
            center.set(qn('w:pos'), str(center_twips))
            tabs.append(center)

            right = OxmlElement('w:tab')
            right.set(qn('w:val'), 'right')
            right.set(qn('w:pos'), str(right_twips))
            tabs.append(right)

            pPr.append(tabs)

        exam_type = self._get_exam_type_title(exam_data, is_answer_key, is_special)
        teacher_name = self._sanitize_text(
            exam_data.get('teacher_name')
            or exam_data.get('created_by')
            or exam_data.get('created_by_name')
            or "N/A"
        )
        subject_name = self._sanitize_text(
            exam_data.get('subject_name')
            or exam_data.get('subject')
            or exam_data.get('module_title')
            or "N/A"
        )

        # One-line identity row: Created by (left) | Exam type (center) | Subject (right)
        p1 = doc.add_paragraph()
        _set_center_right_tabs(p1)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(0)

        run_left = p1.add_run(f"Created by: {teacher_name}")
        run_left.font.size = Pt(10)

        p1.add_run("\t")
        run_center = p1.add_run(exam_type)
        run_center.font.size = Pt(11)
        run_center.font.bold = False
        if is_answer_key:
            run_center.font.color.rgb = RGBColor(255, 0, 0)

        p1.add_run("\t")
        run_right = p1.add_run(f"Subject: {subject_name}")
        run_right.font.size = Pt(10)

        # Second line: exam title centered.
        course_title = self._sanitize_text(exam_data.get('title', 'EXAMINATION'))
        p2 = doc.add_paragraph(course_title.upper())
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.runs[0].font.size = Pt(11)
        p2.runs[0].font.bold = False
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after  = Pt(8)

    def _add_exam_metadata(self, doc, exam_data):
        """Legacy no-op (metadata is now rendered in _add_exam_title)."""
        return
    
    def export_exam(self, exam_data, output_path, include_header=True):
        """
        Export exam to Word - EXACT FORMAT with COMPRESSED SPACING
        """
        self._temp_export_paths = []
        try:
            logger.info(f"Starting exam Word export to {output_path}")
            
            questions = exam_data.get('questions', [])
            
            if not questions:
                logger.error("No questions provided for export")
                return False
            
            doc = Document()
            
            # Set margins
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)  # Reduced margins
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Set COMPRESSED line spacing for entire document
            style = doc.styles['Normal']
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            style.paragraph_format.space_after = Pt(0)  # No default spacing
            style.paragraph_format.space_before = Pt(0)
            
            # 1. HEADER
            if include_header:
                self._add_exam_header(doc, exam_data)
            
            # 2. TITLE
            is_special = exam_data.get('is_special', False)
            self._add_exam_title(doc, exam_data, is_answer_key=False, is_special=is_special)
            
            # 3. GROUP QUESTIONS BY TYPE
            grouped = {}
            for q in questions:
                q_type = q.get('question_type', 'multiple_choice')
                if q_type not in grouped:
                    grouped[q_type] = []
                grouped[q_type].append(q)
            
            # 4. QUESTION TYPE ORDER
            type_order = ['multiple_choice', 'true_false', 'fill_in_blank', 'identification']
            
            # 5. ADD QUESTIONS with compressed spacing
            question_number = 1
            for q_type in type_order:
                if q_type not in grouped:
                    continue
                
                questions_list = grouped[q_type]
                
                if q_type == 'multiple_choice':
                    question_number = self._add_multiple_choice_section(doc, questions_list, question_number)
                elif q_type == 'true_false':
                    question_number = self._add_true_false_section(doc, questions_list, question_number)
                elif q_type == 'fill_in_blank':
                    question_number = self._add_fill_in_blank_section(doc, questions_list, question_number)
                elif q_type == 'identification':
                    question_number = self._add_identification_table(doc, questions_list, question_number)
            
            doc.save(output_path)
            logger.info(f"✅ Exam Word document created: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Error exporting exam to Word: {str(e)}", exc_info=True)
            return False
        finally:
            self._cleanup_temp_export_paths()

    def _add_section_header(self, doc, numeral, instruction):
        """Add a Roman-numeral section header matching the reference docx.
        e.g.  I.  MULTIPLE CHOICE: Choose the letter of the correct answer.
        """
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(4)
        run_num = p.add_run(f"{numeral}.  ")
        run_num.font.bold = True
        run_num.font.size = Pt(11)
        run_instr = p.add_run(instruction)
        run_instr.font.bold = True
        run_instr.font.size = Pt(11)

    def _extract_section_instruction(self, questions, default_instruction):
        """
        Backward-compat cleanup:
        Older generated exams prepend section instructions into the first question
        as: "<instruction>\\n\\n<actual question>".
        Move that instruction to the section header during export.
        """
        if not questions:
            return default_instruction

        for question in questions:
            explicit_instruction = self._sanitize_text(question.get('section_instruction', ''))
            if explicit_instruction:
                return explicit_instruction

        first_question = questions[0]
        first_text = self._sanitize_text(first_question.get('question_text', ''))
        if "\n\n" not in first_text:
            return default_instruction

        maybe_instruction, remaining_text = first_text.split("\n\n", 1)
        maybe_instruction = self._sanitize_text(maybe_instruction)
        remaining_text = self._sanitize_text(remaining_text)

        if not maybe_instruction or not remaining_text:
            return default_instruction

        if len(maybe_instruction) > 240:
            return default_instruction

        first_question['question_text'] = remaining_text
        return maybe_instruction

    @staticmethod
    def _is_mcq_option_correct(option_label, option_text, correct_answer):
        """Return True when label/text corresponds to the saved correct answer."""
        normalized_correct = str(correct_answer or "").strip()
        if not normalized_correct:
            return False

        normalized_label = str(option_label or "").strip().upper()
        normalized_correct_upper = normalized_correct.upper()
        normalized_option = str(option_text or "").strip().lower()

        if normalized_option and normalized_option == normalized_correct.lower():
            return True

        if normalized_correct_upper in {normalized_label, f"{normalized_label}.", f"{normalized_label})"}:
            return True

        if normalized_correct_upper.startswith(f"{normalized_label}.") or normalized_correct_upper.startswith(f"{normalized_label})"):
            return True

        return False

    @staticmethod
    def _remove_table_borders(table):
        """Hide all table borders in Word output."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        tbl = table._tbl
        tbl_pr = tbl.tblPr
        if tbl_pr is None:
            return

        for child in list(tbl_pr):
            if child.tag == qn('w:tblBorders'):
                tbl_pr.remove(child)

        tbl_borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            edge_element = OxmlElement(f'w:{edge}')
            edge_element.set(qn('w:val'), 'nil')
            tbl_borders.append(edge_element)
        tbl_pr.append(tbl_borders)

    def _add_mcq_option_grid(self, doc, options, correct_answer=None):
        """Render MCQ choices in aligned two-column rows: A/C then B/D."""
        option_rows = [
            (('A', 0), ('C', 2)),
            (('B', 1), ('D', 3)),
        ]

        table = doc.add_table(rows=2, cols=2)
        table.autofit = False
        self._remove_table_borders(table)

        for row in table.rows:
            row.cells[0].width = Inches(3.2)
            row.cells[1].width = Inches(3.2)

        for row_index, pair in enumerate(option_rows):
            for col_index, (label, opt_index) in enumerate(pair):
                cell = table.cell(row_index, col_index)
                paragraph = cell.paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0

                if opt_index >= len(options):
                    continue

                option_text = self._sanitize_text(options[opt_index])
                run = paragraph.add_run(f"{label}. {option_text}")
                run.font.size = Pt(10)

                if correct_answer is not None and self._is_mcq_option_correct(label, option_text, correct_answer):
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0, 128, 0)

    def _add_multiple_choice_section(self, doc, questions, start_number=1):
        """Add MCQ - EXACT FORMAT matching reference docx with COMPRESSED SPACING"""
        section_instruction = self._extract_section_instruction(
            questions,
            "Choose the letter of the correct answer."
        )
        self._add_section_header(
            doc, "I",
            f"MULTIPLE CHOICE: {section_instruction}"
        )
        
        question_number = start_number
        for i, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            options = q.get('options', [])
            
            if isinstance(options, str):
                try:
                    options = json.loads(options)
                except:
                    options = []
            
            # Question - JUSTIFIED with compressed spacing
            p_q = doc.add_paragraph(f"{question_number}. {question_text}")
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_q.runs[0].font.size = Pt(10)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(1)  # Minimal spacing
            p_q.paragraph_format.line_spacing = 1.0  # Single spacing

            self._add_question_image(doc, q)
            
            # Options - fixed two-column grid (A/C on row 1, B/D on row 2)
            if len(options) >= 2:
                self._add_mcq_option_grid(doc, options)
            
            # COMPRESSED SPACING between questions - minimal
            if i < len(questions) - 1:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(3)  # Very small space
            question_number += 1
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
        return question_number
    
    def _add_true_false_section(self, doc, questions, start_number=1):
        """Add T/F - EXACT FORMAT matching reference docx with COMPRESSED SPACING"""
        section_instruction = self._extract_section_instruction(
            questions,
            "Read each statement, choose True if the statement is true, otherwise choose False."
        )
        self._add_section_header(
            doc, "II",
            f"TRUE OR FALSE: {section_instruction}"
        )
        
        question_number = start_number
        for _, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            
            p_q = doc.add_paragraph(f"{question_number}. {question_text}")
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_q.runs[0].font.size = Pt(10)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(2)  # Minimal spacing
            p_q.paragraph_format.line_spacing = 1.0

            self._add_question_image(doc, q)
            question_number += 1
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
        return question_number
    
    def _add_fill_in_blank_section(self, doc, questions, start_number=1):
        """Add Fill-in-Blank section matching reference docx with COMPRESSED SPACING"""
        section_instruction = self._extract_section_instruction(
            questions,
            "Write the correct word or phrase on the space provided."
        )
        self._add_section_header(
            doc, "III",
            f"FILL IN THE BLANK: {section_instruction}"
        )
        
        question_number = start_number
        for _, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            
            p_q = doc.add_paragraph(f"{question_number}. {question_text}")
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_q.runs[0].font.size = Pt(10)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(2)
            p_q.paragraph_format.line_spacing = 1.0

            self._add_question_image(doc, q)
            question_number += 1
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
        return question_number
    
    def _add_identification_table(self, doc, questions, start_number=1):
        """Add identification section matching reference docx"""
        if not questions:
            return start_number

        section_instruction = self._extract_section_instruction(
            questions,
            "Identify the term or concept being described. Write your answer on the space provided."
        )

        self._add_section_header(
            doc, "IV",
            f"IDENTIFICATION: {section_instruction}"
        )

        question_number = start_number
        for _, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            p_q = doc.add_paragraph(f"{question_number}. {question_text} _________________")
            p_q.runs[0].font.size = Pt(10)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(3)
            p_q.paragraph_format.line_spacing = 1.0

            self._add_question_image(doc, q)
            question_number += 1

        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
        return question_number
    
    def export_answer_key(self, answer_key_data, output_path, include_header=True):
        """
        Export answer key - EXACT FORMAT with COMPRESSED SPACING
        """
        try:
            logger.info(f"Starting answer key export to {output_path}")
            
            # Handle both possible data structures
            if 'questions' in answer_key_data:
                questions = answer_key_data['questions']
                exam_data = answer_key_data
            elif 'answer_key' in answer_key_data:
                questions = answer_key_data['answer_key'].get('questions', [])
                exam_data = answer_key_data['answer_key']
            else:
                questions = answer_key_data.get('questions', [])
                exam_data = answer_key_data
            
            if not questions:
                logger.error("No questions for answer key")
                return False
            
            doc = Document()
            
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)
            
            # Set COMPRESSED spacing
            style = doc.styles['Normal']
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            style.paragraph_format.space_after = Pt(0)
            style.paragraph_format.space_before = Pt(0)
            
            # HEADER
            if include_header:
                self._add_exam_header(doc, exam_data)
            
            # TITLE with "ANSWER KEY"
            is_special = exam_data.get('is_special', False)
            self._add_exam_title(doc, exam_data, is_answer_key=True, is_special=is_special)
            
            # Group questions
            grouped = {}
            for q in questions:
                q_type = q.get('question_type', 'multiple_choice')
                if q_type not in grouped:
                    grouped[q_type] = []
                grouped[q_type].append(q)
            
            type_order = ['multiple_choice', 'true_false', 'fill_in_blank', 'identification']
            
            # Answers only (no question text)
            self._add_answers_only(doc, questions)
            
            doc.save(output_path)
            logger.info(f"✅ Answer key created: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ Answer key error: {str(e)}", exc_info=True)
            return False

    def _add_answers_only(self, doc, questions):
        """Add answers only, hiding the question text."""
        for i, q in enumerate(questions, 1):
            ans = self._sanitize_text(q.get('correct_answer', ''))
            pts = q.get('points', 1)
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {ans}  ({pts} pt{'s' if pts != 1 else ''})")
            run.bold = True
            run.font.size = Pt(11)
            para.paragraph_format.space_after = Pt(4)
            para.paragraph_format.line_spacing = 1.0
    
    def _add_mcq_with_answers(self, doc, questions):
        """MCQ with answers GREEN/BOLD - COMPRESSED SPACING"""
        self._add_section_header(
            doc, "I",
            "MULTIPLE CHOICE: Choose the letter of the correct answer."
        )
        
        for i, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            options = q.get('options', [])
            correct_answer = self._sanitize_text(q.get('correct_answer', ''))
            
            if isinstance(options, str):
                try:
                    options = json.loads(options)
                except:
                    options = []
            
            p_q = doc.add_paragraph(f"{i+1}. {question_text}")
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_q.runs[0].font.size = Pt(10)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(1)
            p_q.paragraph_format.line_spacing = 1.0
            
            if len(options) >= 2:
                self._add_mcq_option_grid(doc, options, correct_answer=correct_answer)
            
            # COMPRESSED SPACING between questions
            if i < len(questions) - 1:
                spacer = doc.add_paragraph()
                spacer.paragraph_format.space_after = Pt(3)
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
    
    def _add_tf_with_answers(self, doc, questions):
        """T/F with answers [T] or [F] in GREEN/BOLD - COMPRESSED SPACING"""
        self._add_section_header(
            doc, "II",
            "TRUE OR FALSE: Read each statement, choose True if the statement is true, otherwise choose False."
        )
        
        for i, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            correct_answer = q.get('correct_answer', 'True')
            
            p_q = doc.add_paragraph()
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(2)
            p_q.paragraph_format.line_spacing = 1.0
            
            run1 = p_q.add_run(f"{i+1}. {question_text} ")
            run1.font.size = Pt(10)
            
            answer_letter = 'T' if correct_answer == 'True' else 'F'
            run2 = p_q.add_run(f"[{answer_letter}]")
            run2.font.bold = True
            run2.font.color.rgb = RGBColor(0, 128, 0)
            run2.font.size = Pt(10)
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
    
    def _add_fib_with_answers(self, doc, questions):
        """Fill-in-Blank with answers in [brackets] - COMPRESSED SPACING"""
        self._add_section_header(
            doc, "III",
            "FILL IN THE BLANK: Write the correct word or phrase on the space provided."
        )
        
        for i, q in enumerate(questions):
            question_text = self._sanitize_text(q.get('question_text', ''))
            correct_answer = self._sanitize_text(q.get('correct_answer', ''))
            
            answer_filled = question_text.replace('_' * 10, f'[{correct_answer}]')
            
            p_q = doc.add_paragraph(f"{i+1}. {answer_filled}")
            p_q.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p_q.runs[0].font.size = Pt(10)
            p_q.paragraph_format.space_before = Pt(0)
            p_q.paragraph_format.space_after = Pt(2)
            p_q.paragraph_format.line_spacing = 1.0
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
    
    def _add_identification_answers(self, doc, questions):
        """Identification answers in GREEN - COMPRESSED SPACING"""
        if not questions:
            return

        self._add_section_header(
            doc, "IV",
            "IDENTIFICATION ANSWERS:"
        )
        
        for i, q in enumerate(questions):
            p_ans = doc.add_paragraph(f"{i+1}. {self._sanitize_text(q.get('correct_answer', ''))}")
            p_ans.runs[0].font.size = Pt(10)
            p_ans.runs[0].font.color.rgb = RGBColor(0, 128, 0)
            p_ans.paragraph_format.space_before = Pt(0)
            p_ans.paragraph_format.space_after = Pt(2)
            p_ans.paragraph_format.line_spacing = 1.0
        
        # COMPRESSED SPACING after section
        final_spacer = doc.add_paragraph()
        final_spacer.paragraph_format.space_after = Pt(6)
    
    def export_tos(self, tos_data, output_path):
        """
        Export TOS - PRESERVED functionality
        """
        try:
            logger.info(f"Starting TOS export to {output_path}")
            
            if not tos_data:
                logger.error("No TOS data")
                return False
            
            doc = Document()
            
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)
            
            title = doc.add_heading("Table of Specification", 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            exam_title = self._sanitize_text(tos_data.get('exam_title', 'Untitled Exam'))
            doc.add_paragraph(f"Exam: {exam_title}")
            doc.add_paragraph(f"Total Questions: {tos_data.get('total_questions', 0)}")
            doc.add_paragraph(f"Duration: {tos_data.get('duration_minutes', 60)} minutes")
            
            doc.add_paragraph()
            
            # Cognitive Distribution
            cognitive_heading = doc.add_heading('Cognitive Level Distribution', level=1)
            for run in cognitive_heading.runs:
                run.font.color.rgb = RGBColor(37, 99, 235)
            
            cognitive_dist = tos_data.get('cognitive_distribution', {})
            cognitive_pct = tos_data.get('cognitive_percentages', {})
            
            if cognitive_dist:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Cognitive Level'
                hdr_cells[1].text = 'Number of Questions'
                hdr_cells[2].text = 'Percentage'
                
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for level, count in cognitive_dist.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = level.replace('_', ' ').title()
                    row_cells[1].text = str(count)
                    percentage = cognitive_pct.get(level, 0)
                    row_cells[2].text = f"{percentage}%"
            else:
                doc.add_paragraph("No cognitive level data available")
            
            doc.add_paragraph()
            
            # Difficulty Distribution
            difficulty_heading = doc.add_heading('Difficulty Level Distribution', level=1)
            for run in difficulty_heading.runs:
                run.font.color.rgb = RGBColor(16, 185, 129)
            
            difficulty_dist = tos_data.get('difficulty_distribution', {})
            difficulty_pct = tos_data.get('difficulty_percentages', {})
            
            if difficulty_dist:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Light Grid Accent 1'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Difficulty Level'
                hdr_cells[1].text = 'Number of Questions'
                hdr_cells[2].text = 'Percentage'
                
                for cell in hdr_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                for level, count in difficulty_dist.items():
                    row_cells = table.add_row().cells
                    row_cells[0].text = level.title()
                    row_cells[1].text = str(count)
                    percentage = difficulty_pct.get(level, 0)
                    row_cells[2].text = f"{percentage}%"
            else:
                doc.add_paragraph("No difficulty level data available")
            
            # Topic-Cognitive Matrix
            if 'topic_cognitive_matrix' in tos_data and tos_data['topic_cognitive_matrix']:
                doc.add_paragraph()
                
                matrix_heading = doc.add_heading('Topic-Cognitive Level Matrix', level=1)
                for run in matrix_heading.runs:
                    run.font.color.rgb = RGBColor(139, 92, 246)
                
                topics = list(tos_data['topic_cognitive_matrix'].keys())
                cognitive_levels = list(cognitive_dist.keys())
                
                if topics and cognitive_levels:
                    table = doc.add_table(rows=1, cols=len(cognitive_levels) + 1)
                    table.style = 'Light Grid Accent 1'
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Topic'
                    for i, level in enumerate(cognitive_levels):
                        hdr_cells[i + 1].text = level.replace('_', ' ').title()
                    
                    for cell in hdr_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
                    
                    for topic in topics:
                        row_cells = table.add_row().cells
                        row_cells[0].text = self._sanitize_text(topic)
                        for i, level in enumerate(cognitive_levels):
                            count = tos_data['topic_cognitive_matrix'][topic].get(level, 0)
                            row_cells[i + 1].text = str(count)
                else:
                    doc.add_paragraph("No topic-cognitive matrix data")
            
            doc.save(output_path)
            logger.info(f"✅ TOS created: {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"❌ TOS error: {str(e)}", exc_info=True)
            return False
